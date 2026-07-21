from docx import Document


def create_docx(report, filename):

    doc = Document()

    doc.add_heading(
        "AI Spreadsheet Report",
        0
    )

    doc.add_paragraph(

        "Generated: "

        + report["generated_at"]

    )

    doc.add_heading(
        "Business Insights",
        level=1
    )

    for tip in report["insights"]:

        doc.add_paragraph(

            tip,

            style="List Bullet"

        )

    doc.save(filename)

    return filename